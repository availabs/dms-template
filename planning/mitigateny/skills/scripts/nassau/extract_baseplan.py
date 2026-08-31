"""Phase 6 — extract the Nassau base plan's Roles / Participation / county-HOC sources.

Tables, by role:
  0   71x9  jurisdiction x meeting attendance + adoption status  -> Participation, Roles
  1  191x5  the 190-person roster                               -> Roles
  7   10x4  the 9 dated meetings                                 -> Participation
  9/10/11/12/13  hazard identification, reasons, exclusions, probability, ranking -> county HOC
  9x2 boxes      per-hazard Rank/Impact/Frequency/... profiles   -> county HOC

usage: extract_baseplan.py
"""
import collections
import io
import json
import os
import re
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
from annex_lib import cell_lines, row_cell_objs, table_grid       # noqa: E402
from docx import Document                                         # noqa: E402

BASE = ('references/mny-transcribe/Nassau/Base Plan and Appendices/'
        'Nassau County_HMP_Base_Plan_12.16.20.docx')
CTX = 'references/mny-transcribe/Nassau/context/'
OUT = CTX + 'extracted/baseplan.json'

MONTHS = ('January February March April May June July August September October November '
          'December').split()
DATE = re.compile(r'(' + '|'.join(MONTHS) + r')\s+(\d{1,2})(?:\s*and\s*(\d{1,2}))?,\s*(\d{4})')


def parse_dates(s):
    """Return every explicit date in a cell, expanding "February 19 and 20, 2020" to two."""
    out = []
    for m in DATE.finditer(s or ''):
        mon, d1, d2, yr = m.groups()
        mi = MONTHS.index(mon) + 1
        out.append(f'{yr}-{mi:02d}-{int(d1):02d}')
        if d2:
            out.append(f'{yr}-{mi:02d}-{int(d2):02d}')
    return out


def main():
    doc = Document(BASE)
    grids = [table_grid(t)[0] for t in doc.tables]
    res = {'source': os.path.basename(BASE), 'warnings': []}

    # ---- table 7: the dated meetings -------------------------------------
    meetings = []
    g = grids[7]
    for row in g[1:]:
        if not row or not row[0].strip():
            continue
        name, raw_date, desc = row[0].strip(), row[1].strip(), row[2].strip()
        part = row[3].strip() if len(row) > 3 else ''
        dates = parse_dates(raw_date)
        is_range = '–' in raw_date or '-' in raw_date
        if not dates:
            res['warnings'].append(f'meeting "{name}": no parseable date in "{raw_date}"')
        rows_for = dates if dates else ['']
        for i, d in enumerate(rows_for):
            meetings.append({
                'meeting_name': name, 'date': d, 'date_verbatim': raw_date,
                'narrative': desc, 'participation': part,
                'meeting_unique_id': re.sub(r'\W+', '_', name).strip('_').lower(),
                'row_of': f'{i + 1}/{len(rows_for)}',
                'is_date_range': bool(is_range and len(rows_for) > 1),
                'format': ('Virtual' if 'webinar' in name.lower()
                           else 'Phone Call' if 'call' in name.lower()
                           else 'In-Person' if ('workshop' in name.lower()
                                               or 'meeting' in name.lower()) else ''),
            })
    res['meetings'] = meetings

    # ---- table 0: attendance matrix -------------------------------------
    g = grids[0]
    hdr = [c.strip() for c in g[0]]
    meeting_cols = hdr[1:-1]
    attendance, adoption = [], {}
    for row in g[1:]:
        if not row or not row[0].strip():
            continue
        juris = row[0].strip()
        adoption[juris] = row[-1].strip()
        for i, col in enumerate(meeting_cols, start=1):
            if i < len(row) and row[i].strip():
                attendance.append({'jurisdiction_label': juris, 'meeting_column': col,
                                   'mark': row[i].strip()})
    res['attendance'] = attendance
    res['adoption_status'] = adoption
    res['meeting_columns'] = meeting_cols

    # ---- table 1: the roster --------------------------------------------
    g = grids[1]
    roster = []
    for row in g[1:]:
        if len(row) < 5 or not (row[1].strip() or row[2].strip()):
            continue
        roster.append({'organization': row[0].strip(),
                       'name': f'{row[1].strip()} {row[2].strip()}'.strip(),
                       'first_name': row[1].strip(), 'last_name': row[2].strip(),
                       'title': row[3].strip(), 'core_planning_group': row[4].strip()})
    res['roster'] = roster

    # ---- county hazard identification / exclusion / ranking -------------
    res['hazards_identified'] = [
        {'hazard': r[0].strip(), 'reason_for_identification': r[1].strip(),
         'connection_to_2014_plan': r[2].strip() if len(r) > 2 else ''}
        for r in grids[10][1:] if r and r[0].strip()]
    # The six names are separate PARAGRAPHS inside two merged cells. cell_text() joins
    # paragraphs with a single space, which destroys the boundaries -- read lines instead.
    res['hazards_not_profiled'] = [ln for c in row_cell_objs(doc.tables[11].rows[0])
                                   for ln in cell_lines(c) if ln.strip()]
    res['probability_categories'] = [
        {'category': r[0].strip(), 'hazards': [h.strip() for h in r[1].split(',') if h.strip()]}
        for r in grids[12][1:] if r and r[0].strip()]
    ranking = []
    for r in grids[13][1:]:
        for a, b, c in ((0, 1, 2), (3, 4, 5)):
            if len(r) > c and r[b].strip():
                ranking.append({'rank': r[a].strip(), 'hazard': r[b].strip(),
                                'hazard_rank': r[c].strip()})
    res['hazard_ranking'] = ranking

    # ---- the per-hazard 9x2 profile boxes -------------------------------
    boxes = []
    for i, g in enumerate(grids):
        # NOTE: row 0 is a MERGED banner carrying the hazard name, so after merged-cell
        # dedupe it has ONE cell, not two. Requiring len(g[0]) >= 2 finds nothing.
        if 8 <= len(g) <= 10 and len(g) > 1 and len(g[1]) >= 2 and g[1][0].strip() == 'Rank':
            box = {'table_index': i, 'hazard': g[0][0].strip()}
            for row in g[1:]:
                if len(row) >= 2 and row[0].strip():
                    box[row[0].strip()] = row[1].strip()
            boxes.append(box)
    res['hazard_profile_boxes'] = boxes

    io.open(OUT, 'w', encoding='utf-8').write(json.dumps(res, indent=1, ensure_ascii=False))
    print(f'wrote {OUT}')
    print(f'  meetings (rows)          {len(meetings)}  from {len({m["meeting_name"] for m in meetings})} meetings')
    for m in meetings:
        flag = '  [range bound]' if m['is_date_range'] else ''
        print(f'      {m["date"] or "??":10s} {m["row_of"]:4s} {m["meeting_name"][:52]:54s}'
              f'{m["format"]:11s}{flag}')
    print(f'  attendance marks         {len(attendance)}')
    print(f'  adoption status          {dict(collections.Counter(adoption.values()))}')
    print(f'  roster people            {len(roster)}   CPG={dict(collections.Counter(r["core_planning_group"] for r in roster))}')
    print(f'  hazards identified       {len(res["hazards_identified"])}')
    print(f'  hazards NOT profiled     {res["hazards_not_profiled"]}')
    print(f'  probability categories   {[(p["category"], len(p["hazards"])) for p in res["probability_categories"]]}')
    print(f'  hazard ranking rows      {len(ranking)}')
    print(f'  9x2 profile boxes        {len(boxes)}  -> {[b["hazard"][:26] for b in boxes]}')
    if res['warnings']:
        print('  warnings:', res['warnings'])


if __name__ == '__main__':
    main()
